import PySimpleGUI as sg
import aspose.words as aw
from docx import Document
import os
import docx2pdf as pd

sg.theme('SystemDefaultForReal')
def replaceh(document1):
	document=Document(document1)
	document.paragraphs[0].text=" "

	for section in document.sections:
		footer = section.footer
		footer.paragraphs[0].text  = " "
	document.save("c1.docx")
	os.remove("c.docx")

layout=[[sg.Text("Telephone:	"), sg.In(key="tel_num")],
		[sg.Text("ONT Serial No.:	"), sg.In(key="ont_ser_num")],
		[sg.Text("RTOM Area:	"), sg.In(key="rtom_area")],
		[sg.Text("Contractor:	"), sg.In(key="cont")],
		[sg.Text("Name:		"), sg.In(key="name")],
		[sg.Text("Designation:	"), sg.In(key="desig")],
		[sg.Text("Mobile Number:	"), sg.In(key="num")],
		[sg.Button("Save a Copy"), sg.Button("Print")]
		]
window=sg.Window("Doc Filler", layout, icon="icon.ico")

while True:
	event, values=window.read()
	if event==sg.WIN_CLOSED:
		try:
			os.remove("c1.docx")
		except:
			pass
		break
	if event=="Print":
		tel=""
		for i in values["tel_num"]:
			tel=tel+i+"       "

		doc=aw.Document("doc.docx")
		doc.range.replace("Fs_name", values["name"], aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
		doc.range.replace("Fs_desig", values["desig"], aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
		doc.range.replace("Fs_num", values["num"], aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
		doc.range.replace("Ont_ser_num", values["ont_ser_num"], aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
		doc.range.replace("Rtom_area", values["rtom_area"], aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
		doc.range.replace("Cont name", values["cont"], aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
		doc.range.replace("Tel_num", tel, aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
		doc.save("c.docx")

		replaceh("c.docx")

		os.startfile("c1.docx", "print")

	if event=="Save a Copy":
		tel=""
		for i in values["tel_num"]:
			tel=tel+i+"       "

		doc=aw.Document("doc.docx")
		doc.range.replace("Fs_name", values["name"], aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
		doc.range.replace("Fs_desig", values["desig"], aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
		doc.range.replace("Fs_num", values["num"], aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
		doc.range.replace("Ont_ser_num", values["ont_ser_num"], aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
		doc.range.replace("Rtom_area", values["rtom_area"], aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
		doc.range.replace("Cont name", values["cont"], aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
		doc.range.replace("Tel_num", tel, aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
		doc.save("c.docx")

		replaceh("c.docx")

		pd.convert("c1.docx", "Doc.pdf")

window.close()